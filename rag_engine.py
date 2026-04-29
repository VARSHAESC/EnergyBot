# -*- coding: utf-8 -*-
"""
rag_engine.py — Multi-Utility Hybrid Engine (Gas | Strom | Wasser)
Strictly Offline Version using Gemma 3 via Ollama.

FIXED: Risk assessment accuracy, engineering standards alignment,
       LLM prompt completeness, and dataframe-based risk explanations.
"""

from __future__ import annotations

# --- NumPy 2.x compatibility shim ---
import numpy as _np
if not hasattr(_np, "float_"):   _np.float_ = _np.float64
if not hasattr(_np, "int_"):     _np.int_ = int
if not hasattr(_np, "bool_"):    _np.bool_ = bool
if not hasattr(_np, "object_"): _np.object_ = object
# ------------------------------------

import os
import re
import io
from typing import List, Dict, Any, Optional, Sequence, Set

import pandas as pd
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv
import requests
import json
import docx

from geo_utils import (
    load_excel, CSV_FILES, ALL_UTILITIES, MATERIAL_LIFESPAN,
    get_utility_df, get_unified_df, RISK_LADDER, CURRENT_YEAR
)

load_dotenv()

# ─────────────── Config ───────────────
OFFLINE: bool = os.getenv("OFFLINE_MODE", "false").lower() == "true"
PERSIST_DIR: str = "./chroma_db"
EMBED_MODEL_NAME: str = os.getenv(
    "EMBED_MODEL_NAME",
    "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
)
RETURN_ALL_MAX: int = 100_000

# ─────────────── Utilities ───────────────
def _safe(v: Any) -> str:
    return "" if pd.isna(v) else str(v)

def row_to_paragraph(row: Dict[str, Any], utility: str = "") -> str:
    general_keys = ["Gemeinde", "Postleitzahl", "Straße", "Hausnummer", "Objekt-ID_Global"]
    util_specific = {k: v for k, v in row.items() if k not in general_keys and k != "Sparte"}
    parts = []
    if utility: parts.append(f"VERSORGUNGSART: {utility}")
    gen = [f"{k}: {_safe(row.get(k))}" for k in general_keys if row.get(k) and pd.notna(row.get(k))]
    if gen: parts.append("ALLGEMEINE OBJEKTDATEN: " + ", ".join(gen))
    spec = [f"{k}: {_safe(v)}" for k, v in util_specific.items() if pd.notna(v) and str(v).strip() not in ("", "nan")]
    if spec: parts.append(f"DATEN ZUM NETZANSCHLUSS {utility}: " + ", ".join(spec))
    return " | ".join(parts) + "."


# ─────────────── Engineering Standards ────────────────────────────────────────
# FIX: Build ENGINEERING_STANDARDS dynamically from RISK_LADDER in geo_utils.py
# so there is ONE single source of truth. No more hand-typed thresholds that
# can drift out of sync with the actual classification logic.
# ─────────────────────────────────────────────────────────────────────────────
def _build_engineering_standards() -> str:
    """
    Dynamically generates the Engineering Standards text from the RISK_LADDER
    defined in geo_utils.py. This guarantees the LLM always sees the same
    thresholds that the classification engine uses — no duplication, no drift.
    """
    lines = [
        "TECHNICAL LIFE TABLES — RISK CLASSIFICATION (Risikobewertung):",
        "=" * 60,
        "Risk levels are assigned per material based on installation age (Alter = CURRENT_YEAR - Einbaujahr).",
        "",
        "Rule:  age <= gut_threshold  →  Niedrig (Low Risk)",
        "       age <= mittel_threshold →  Mittel  (Medium Risk)",
        "       age >  mittel_threshold →  Hoch    (High Risk)",
        "",
    ]
    for sparte, profiles in RISK_LADDER.items():
        lines.append(f"── {sparte.upper()} ──")
        seen = set()
        for p in profiles:
            mat = p["material"]
            if mat in seen:
                continue
            seen.add(mat)
            lines.append(
                f"  {mat:30s}  Niedrig: 0–{p['gut']}J  |  Mittel: {p['gut']+1}–{p['mittel']}J  "
                f"|  Hoch: >{p['mittel']}J  |  TechNutzungsdauer: {p['life']}J"
            )
        lines.append("")

    # Add materials that use the generic MATERIAL_LIFESPAN fallback
    lines.append("── FALLBACK (materials not in sparte-specific ladder) ──")
    for mat, life in MATERIAL_LIFESPAN.items():
        lines.append(
            f"  {mat:30s}  Niedrig: 0–{int(life*0.65)}J  |  Mittel: {int(life*0.65)+1}–{int(life*0.85)}J  "
            f"|  Hoch: >{int(life*0.85)}J  |  TechNutzungsdauer: {life}J"
        )
    lines.append("")
    lines.append(f"Reference year for age calculation: {CURRENT_YEAR}")
    return "\n".join(lines)

ENGINEERING_STANDARDS: str = _build_engineering_standards()


def _risk_explanation(row: pd.Series) -> str:
    """
    Returns a human-readable explanation of WHY a connection has its risk level,
    referencing the exact thresholds from RISK_LADDER.
    Used both in dataframe answers and as LLM context.
    """
    sparte   = str(row.get("Sparte", ""))
    material = str(row.get("Werkstoff", "Unbekannt")).strip()
    alter    = row.get("Alter", 0)
    risiko   = str(row.get("Risiko", "Unbekannt"))

    if not alter or pd.isna(alter):
        return f"Risiko: {risiko} (Einbaudatum unbekannt — kein Alter berechenbar)"

    alter = float(alter)

    # Look up the profile that was used
    profile = None
    if sparte in RISK_LADDER:
        mat_lower = material.lower()
        for p in RISK_LADDER[sparte]:
            if p["material"].lower() in mat_lower:
                profile = p
                break

    if profile:
        gut    = profile["gut"]
        mittel = profile["mittel"]
        life   = profile["life"]
        if alter <= gut:
            return (
                f"Risiko: **Niedrig** — Material '{material}' ({sparte}), Alter {int(alter)}J. "
                f"Schwellenwert Niedrig: ≤{gut}J, Mittel: {gut+1}–{mittel}J, Hoch: >{mittel}J. "
                f"TechNutzungsdauer: {life}J."
            )
        elif alter <= mittel:
            remaining = life - alter
            return (
                f"Risiko: **Mittel** — Material '{material}' ({sparte}), Alter {int(alter)}J. "
                f"Schwellenwert Niedrig: ≤{gut}J, Mittel: {gut+1}–{mittel}J, Hoch: >{mittel}J. "
                f"Noch ca. {int(remaining)}J bis Ende TechNutzungsdauer ({life}J)."
            )
        else:
            overdue = alter - life
            return (
                f"Risiko: **Hoch** ⚠️ — Material '{material}' ({sparte}), Alter {int(alter)}J. "
                f"Schwellenwert Hoch: >{mittel}J (überschritten). "
                f"TechNutzungsdauer ({life}J) um {int(overdue)}J überschritten — Erneuerung dringend!"
            )
    else:
        # Fallback lifespan
        life = MATERIAL_LIFESPAN.get(material, 50)
        pct  = alter / life
        status = "Hoch" if pct >= 0.85 else "Mittel" if pct >= 0.65 else "Niedrig"
        return (
            f"Risiko: **{status}** — Material '{material}' ({sparte}), Alter {int(alter)}J. "
            f"Fallback-TechNutzungsdauer: {life}J ({int(pct*100)}% verbraucht)."
        )


# ─────────────── Vector Store ───────────────
class VectorStore:
    def __init__(self, persist_dir: str = PERSIST_DIR, name: str = "energy_kb", embed_model: str = EMBED_MODEL_NAME):
        self.client = chromadb.PersistentClient(path=persist_dir, settings=Settings(anonymized_telemetry=False))
        try:
            self.col = self.client.get_collection(name)
        except Exception:
            self.col = self.client.create_collection(name, metadata={"embed_model": embed_model})
        meta = self.col.metadata or {}
        if meta.get("embed_model") != embed_model:
            self.client.delete_collection(name)
            self.col = self.client.create_collection(name, metadata={"embed_model": embed_model})

    def reset(self, metadata: Optional[Dict[str, Any]] = None):
        name = self.col.name
        self.client.delete_collection(name)
        self.col = self.client.create_collection(name, metadata=metadata)

    def count(self) -> int:
        try: return self.col.count()
        except: return 0

    def add(self, ids, embeddings, metadatas, documents):
        self.col.add(ids=ids, embeddings=embeddings, metadatas=metadatas, documents=documents)

    def query(self, query_embeddings, top_k: int = 5):
        cnt = self.count()
        if cnt == 0: return {"metadatas": [[]], "documents": [[]], "distances": [[]]}
        return self.col.query(query_embeddings=query_embeddings, n_results=min(top_k, cnt))

class Embedder:
    def __init__(self, model_name: str = EMBED_MODEL_NAME):
        import os
        os.environ["TQDM_DISABLE"] = "1"
        os.environ["HF_HUB_DISABLE_PROGRESS_BARS"] = "1"
        self.model = SentenceTransformer(model_name)
    def embed(self, texts: List[str], batch_size: int = 64) -> List[List[float]]:
        return self.model.encode(texts, batch_size=batch_size, show_progress_bar=False).tolist()


# ─────────────── Main Engine ───────────────
class EnergyRAG:
    def __init__(self, persist_dir: str = PERSIST_DIR, embed_model: str = EMBED_MODEL_NAME):
        self.vs = VectorStore(persist_dir, "energy_kb_multi", embed_model)
        self.embedder = Embedder(embed_model)

        # --- Provider-Agnostic LLM Configuration ---
        self.azure_openai_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        self.azure_openai_key = os.getenv("AZURE_OPENAI_KEY")

        self.llm_api_key = self.azure_openai_key or os.getenv("LLM_API_KEY", os.getenv("GROQ_API_KEY"))
        self.llm_model = os.getenv("LLM_MODEL_NAME", "gpt-5-chat")
        self.llm_base_url = self.azure_openai_endpoint or os.getenv("LLM_BASE_URL", "https://api.groq.com/openai/v1").rstrip("/")

        self.unified_df = get_unified_df()
        self.stats_summary = ""
        self.refresh_stats_manual()  # Builds self.reference_manual


    def _route_query(self, question: str) -> str:
        q = question.lower()

        if any(x in q for x in ["update", "ändern", "setze", "change", "fix", "modify", "korrigier"]):
            return "AGENT"

        if any(x in q for x in ["risiko", "risk", "warum", "why", "gefahr", "kritisch"]):
            return "RISK"

        if any(x in q for x in ["wie viele", "count", "anzahl", "total", "wieviel"]):
            return "COUNT"

        if any(x in q for x in ["liste", "tabelle", "alle", "list", "übersicht"]):
            return "TABLE"

        if any(x in q for x in ["karte", "map", "zeige", "öffne"]):
            return "MAP"

        if any(x in q for x in ["erneuer", "fällig", "renewal", "überfällig"]):
            return "RENEWAL"

        if any(x in q for x in ["bündeln", "bundle", "zusammenfassen", "koordinier"]):
            return "BUNDLE"

        if any(x in q for x in ["dokument", "plan", "protokoll", "akte", "unterlage"]):
            return "DOCS"

        if any(x in q for x in ["muster", "pattern", "auffällig", "erkenntnisse"]):
            return "PATTERN"

        if any(x in q for x in ["störung", "ausfall", "korrelation", "zusammenhang"]):
            return "CORRELATION"

        if any(x in q for x in ["wärmepumpe", "heat pump", "ladeinfrastruktur", "eignung"]):
            return "SUITABILITY"

        return "LLM"



    def apply_update(self, customer_id: str, field_name: str, new_value: str, utility: str) -> Dict[str, Any]:
        """
        Actually updates the dataframe and persists it back to Excel.
        """
        try:
            from geo_utils import EXCEL_FILE, load_excel
            
            # Update the global cache Excel file directly
            raw_df = load_excel()
            mask = raw_df["Kundennummer"].astype(str) == str(customer_id)
            if not mask.any():
                return {"ok": False, "msg": f"Kunde {customer_id} nicht gefunden."}

            # Apply update
            raw_df.loc[mask, field_name] = new_value

            # Save back to Excel
            if EXCEL_FILE and os.path.exists(EXCEL_FILE):
                raw_df.to_excel(EXCEL_FILE, index=False)
            
            return {
                "ok": True,
                "msg": f"{field_name} erfolgreich auf '{new_value}' gesetzt für Kunde {customer_id}."
            }

        except Exception as e:
            return {"ok": False, "msg": str(e)}


    def _load_reference_manual(self) -> str:
        """Loads the Word reference manual for system prompt context."""
        doc_path = os.path.join("excel_data", "Hausanschluss_KI_Referenzhandbuch.docx")
        if not os.path.exists(doc_path):
            return "Reference Manual not found."
        try:
            doc = docx.Document(doc_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except:
            return "Error loading Reference Manual."

    def check_llm_status(self) -> Dict[str, Any]:
        if not self.llm_api_key:
            return {"ok": False, "msg": "API Key fehlt."}
        try:
            if self.azure_openai_endpoint:
                headers = {"api-key": self.llm_api_key}
                # Minimal call to check status
                resp = requests.post(f"{self.llm_base_url}", headers=headers, json={"messages":[{"role":"user","content":"hi"}],"max_tokens":1}, timeout=5)
            else:
                headers = {"Authorization": f"Bearer {self.llm_api_key}"}
                resp = requests.get(f"{self.llm_base_url}/models", headers=headers, timeout=5)

            if resp.status_code in [200, 400]: # 400 is often returned by Azure for malformed/test reqs but shows connectivity
                provider_name = "Azure OpenAI" if self.azure_openai_endpoint else "Online Provider"
                return {"ok": True, "msg": f"{provider_name}: {self.llm_model}"}
            return {"ok": False, "msg": f"LLM Error: {resp.status_code}"}
        except Exception as e:
            return {"ok": False, "msg": f"Verbindungsfehler: {str(e)[:50]}"}

    def refresh_stats_manual(self):
        """
        Builds a dynamic stats summary AND the full engineering standards for the LLM.

        FIX: The old code truncated reference_manual to 1500 chars in the LLM system
        prompt, which silently dropped risk thresholds. Now we store standards and
        stats separately and inject them in full into every prompt.
        """
        if self.unified_df.empty:
            self.unified_df = get_unified_df()
        df = self.unified_df

        total = len(df)
        risks  = df["Risiko"].value_counts().to_dict() if not df.empty else {}
        sparten = df["Sparte"].value_counts().to_dict() if not df.empty else {}

        # Per-sparte risk breakdown — gives LLM concrete numbers to cite
        risk_by_sparte_lines = []
        if not df.empty:
            for sp in df["Sparte"].unique():
                sp_df = df[df["Sparte"] == sp]
                r = sp_df["Risiko"].value_counts().to_dict()
                risk_by_sparte_lines.append(
                    f"  {sp}: Hoch={r.get('Hoch',0)}, Mittel={r.get('Mittel',0)}, "
                    f"Niedrig={r.get('Niedrig',0)}, Unbekannt={r.get('Unbekannt',0)}"
                )

        # Renewal stats
        renewal_overdue, renewal_10yr = 0, 0
        if "Erneuerung_empfohlen_bis" in df.columns:
            ren_num = pd.to_numeric(df["Erneuerung_empfohlen_bis"], errors="coerce")
            renewal_overdue = int((ren_num < CURRENT_YEAR).sum())
            renewal_10yr    = int(((ren_num >= CURRENT_YEAR) & (ren_num <= CURRENT_YEAR + 10)).sum())

        # Document / inspection gaps
        missing_docs_cnt = int((df["Dokumente"] == "Lückenhaft").sum()) if "Dokumente" in df.columns else 0
        no_insp_cnt      = int(df["missing_inspection"].sum()) if "missing_inspection" in df.columns else 0

        # Average age
        avg_age = round(df["Alter"].mean(), 1) if "Alter" in df.columns and not df["Alter"].isna().all() else "?"

        self.stats_summary = (
            f"CURRENT INFRASTRUCTURE SUMMARY (as of {CURRENT_YEAR}):\n"
            f"- TOTAL CONNECTIONS: {total}\n"
            f"- DISTRIBUTION: {', '.join([f'{k}: {v}' for k, v in sparten.items()])}\n"
            f"- AVERAGE AGE: {avg_age} years\n"
            f"- OVERALL RISK LEVELS: {', '.join([f'{k}: {v}' for k, v in risks.items()])}\n"
            f"- RISK BY SPARTE:\n" + "\n".join(risk_by_sparte_lines) + "\n"
            f"- RENEWAL OVERDUE (before {CURRENT_YEAR}): {renewal_overdue}\n"
            f"- RENEWAL DUE WITHIN 10 YEARS (by {CURRENT_YEAR+10}): {renewal_10yr}\n"
            f"- MISSING DOCUMENTS (Lückenhaft): {missing_docs_cnt}\n"
            f"- NO INSPECTION DATE: {no_insp_cnt}\n"
        )

        # FIX: reference_manual now contains FULL engineering standards (no truncation).
        # The LLM system prompt will inject this directly.
        self.reference_manual = (
            self.stats_summary
            + "\n\n"
            + ENGINEERING_STANDARDS
        )

        # Append Word reference doc if present
        ref_path = "excel_data/Hausanschluss_KI_Referenzhandbuch.docx"
        if os.path.exists(ref_path):
            try:
                import docx as _docx
                doc = _docx.Document(ref_path)
                self.reference_manual += "\n" + "\n".join([p.text for p in doc.paragraphs])
            except:
                pass

    def init_or_refresh_kb(self, utility: Optional[str] = None, reset: bool = False) -> int:
        self.unified_df = get_unified_df()
        self.refresh_stats_manual()
        utils = [utility] if utility else ALL_UTILITIES

        if reset:
            old_meta = self.vs.col.metadata
            self.vs.reset(metadata=old_meta)

        total = 0
        for util in utils:
            df = get_utility_df(util)
            if df.empty: continue

            docs, ids, metas = [], [], []
            for i, row in df.iterrows():
                row_id = f"{util}_{row.get('Datensatz', i)}"
                para = row_to_paragraph(row.to_dict(), utility=util)
                if not para: continue
                docs.append(para)
                ids.append(row_id)
                metas.append({
                    "utility": util,
                    "id": str(row.get("Kundennummer", "")),
                    "name": str(row.get("Kundenname", ""))
                })

            if docs:
                batch_size = 500
                for j in range(0, len(docs), batch_size):
                    b_docs  = docs[j:j+batch_size]
                    b_ids   = ids[j:j+batch_size]
                    b_metas = metas[j:j+batch_size]
                    b_embeddings = self.embedder.embed(b_docs)
                    self.vs.add(ids=b_ids, embeddings=b_embeddings, metadatas=b_metas, documents=b_docs)
                    total += len(b_docs)

        return total

    def transcribe_audio(self, audio_bytes: bytes) -> Dict[str, Any]:
        if not self.llm_api_key: return {"ok": False, "text": "API Key fehlt."}
        if not audio_bytes or len(audio_bytes) < 100:
            return {"ok": False, "text": "Audio-Daten zu kurz oder leer."}
        try:
            headers = {"Authorization": f"Bearer {self.llm_api_key}"}
            files = {
                "file": ("audio.webm", io.BytesIO(audio_bytes), "audio/webm"),
                "model": (None, os.getenv("WHISPER_MODEL", "whisper-large-v3")),
            }
            resp = requests.post(
                f"{self.llm_base_url}/audio/transcriptions",
                headers=headers, files=files, timeout=30
            )
            if resp.status_code == 200:
                t = resp.json().get("text", "")
                return {"ok": True, "text": t}
            return {"ok": False, "text": f"LLM Error: {resp.status_code} - {resp.text}"}
        except Exception as e:
            return {"ok": False, "text": f"Verbindungsfehler: {str(e)}"}

    # ─────────────────────────────────────────────────────────────────────────
    # FAST DATAFRAME ENGINE
    # FIX: All risk-related branches now use _risk_explanation() to produce
    #      accurate, threshold-backed answers instead of generic counts.
    # ─────────────────────────────────────────────────────────────────────────
    def _try_dataframe_answer(self, question: str) -> Optional[Dict[str, Any]]:
        ql = (question or "").lower()

        # 0. Bypass for update commands → must go to Agentic Engine
        update_keywords = ["update", "ändern", "setze", "change", "aktualisier", "korrigier", "fix", "put", "schreib"]
        if any(x in ql for x in update_keywords):
            return None

        self.unified_df = get_unified_df()

        # ── 1. RISK DETAIL for a specific connection / ID ─────────────────────
        # Detect when user asks about risk for a specific customer / connection
        risk_keywords = [
            "risiko", "risk", "gefahr", "danger", "sicher", "safe", "kritisch",
            "critical", "warum", "why", "grund", "reason", "erklär", "explain",
            "bewertung", "assessment", "status", "level", "stufe"
        ]
        id_match = re.search(r'(\d+)', ql)

        if id_match and any(x in ql for x in risk_keywords):
            sid = id_match.group(1)
            search_df = self.unified_df.copy()

            target_sparte = None
            if "gas" in ql:   target_sparte = "Gas"
            elif "wasser" in ql: target_sparte = "Wasser"
            if target_sparte:
                search_df = search_df[search_df["Sparte"] == target_sparte]

            matches = pd.DataFrame()
            for col in ["Kundennummer", "Objekt_ID", "Objekt-ID", "Objekt-ID_Global"]:
                if col in search_df.columns:
                    mask = search_df[col].astype(str).str.contains(rf'\b{sid}\b', regex=True, na=False)
                    if mask.any():
                        matches = search_df[mask]
                        break

            if not matches.empty:
                res = matches.iloc[0]
                explanation = _risk_explanation(res)
                lines = [
                    f"🔍 **Risikobewertung: {res.get('Sparte','')} — {res.get('Kundennummer', sid)}**",
                    f"📍 {res.get('Straße','?')} {res.get('Hausnummer','')}",
                    "",
                    explanation,
                    "",
                    f"🏗️ Material: `{res.get('Werkstoff','?')}` | Dimension: `{res.get('Dimension','?')}`",
                    f"📅 Einbaujahr: `{int(res['Einbaujahr']) if pd.notna(res.get('Einbaujahr')) else 'unbekannt'}`"
                    f" | Alter: `{int(res.get('Alter',0))}J`",
                    f"🔄 Erneuerung empfohlen bis: `{res.get('Erneuerung_empfohlen_bis','?')}`",
                ]
                return {"answer": "\n".join(lines), "hits": [], "model_used": "Risk-Detail-Engine-v2", "switched": True}

        # ── 2. General ID Lookup (no risk keywords — show full record) ────────
        if id_match and not any(x in ql for x in [
            "älter", "jahre", "vor", "nach", "count", "anzahl", "old", "older",
            "years", "more", "over", "über", "list", "table", "tabelle", "alle", "all"
        ]):
            sid = id_match.group(1)
            search_df = self.unified_df.copy()

            target_sparte = None
            if "gas" in ql:    target_sparte = "Gas"
            elif "wasser" in ql: target_sparte = "Wasser"
            if target_sparte:
                search_df = search_df[search_df["Sparte"] == target_sparte]

            matches = pd.DataFrame()
            for col in ["Kundennummer", "Objekt_ID", "Objekt-ID", "Objekt-ID_Global"]:
                if col in search_df.columns:
                    mask = search_df[col].astype(str).str.contains(rf'\b{sid}\b', regex=True, na=False)
                    if mask.any():
                        matches = search_df[mask]
                        break

            if not matches.empty:
                res = matches.iloc[0]
                lines = [
                    f"✅ **Datensatz gefunden: {res.get('Sparte', '')} — {res.get('Kundennummer', sid)}**", ""
                ]

                if any(x in ql for x in ["material", "werkstoff", "type", "art"]):
                    lines.append(f"🏗️ **Material:** {res.get('Werkstoff', 'n/a')} ({res.get('Dimension', '')})")
                if any(x in ql for x in ["strasse", "straße", "hausnummer", "address", "location"]):
                    lines.append(f"📍 **Adresse:** {res.get('Straße', 'n/a')} {res.get('Hausnummer', '')}")
                    lines.append(f"🏙️ **Ort:** {res.get('Postleitzahl', '')} {res.get('Gemeinde', '')}")
                if any(x in ql for x in ["alter", "age", "year", "einbau"]):
                    lines.append(f"⏳ **Alter:** {int(res.get('Alter', 0))} Jahre (Einbau: {res.get('Einbaujahr', 'unbekannt')})")

                # Default: show all key fields
                if len(lines) <= 2:
                    lines.append(f"📍 **Adresse:** {res.get('Straße', 'n/a')} {res.get('Hausnummer', '')}")
                    lines.append(f"🏗️ **Material:** {res.get('Werkstoff', 'n/a')}")
                    lines.append(f"⚠️ **Risiko:** {res.get('Risiko', 'n/a')}")
                    lines.append(f"⏳ **Alter:** {int(res.get('Alter', 0))} Jahre")
                    lines.append(f"🔄 **Erneuerung bis:** {res.get('Erneuerung_empfohlen_bis', '?')}")

                # Map navigation
                if any(x in ql for x in ["karte", "map", "zeige", "view", "show"]):
                    if pd.notna(res.get("lat")) and pd.notna(res.get("lon")):
                        return {
                            "answer": (
                                f"📍 **Navigation zur Karte gestartet...**\n"
                                f"Ich zeige Ihnen den Anschluss `{res.get('Kundennummer', sid)}` "
                                f"({res.get('Sparte', '')}) in der "
                                f"`{res.get('Straße', '')} {res.get('Hausnummer', '')}` auf der Karte."
                            ),
                            "hits": [], "model_used": "Navigation-Engine-v1", "switched": True,
                            "pending_action": {
                                "type": "navigate_map",
                                "args": {
                                    "customer_id": str(res.get("Kundennummer", sid)),
                                    "lat": float(res["lat"]),
                                    "lon": float(res["lon"])
                                }
                            }
                        }
                    else:
                        return {
                            "answer": f"⚠️ Kunde `{sid}` gefunden, aber keine Koordinaten verfügbar.",
                            "hits": [], "model_used": "Navigation-Engine-v1", "switched": True
                        }

                dl_requested = any(x in ql for x in ["excel", "csv", "tabelle", "table", "format", "tabular"])
                resp = {"answer": "\n".join(lines), "hits": [], "model_used": "Direct-ID-Engine-v2", "switched": True}
                if dl_requested:
                    resp["download_data"] = matches.to_csv(index=False).encode('utf-8-sig')
                return resp

        # ── 3. RISK OVERVIEW — "how many are high risk / what is the risk?" ──
        # FIX: When user asks a general risk question (no specific ID), give a
        # full material-level breakdown so the answer is educational, not just a count.
        if any(x in ql for x in risk_keywords) and not id_match:
            target_sparte = None
            if "gas" in ql:    target_sparte = "Gas"
            elif "wasser" in ql: target_sparte = "Wasser"

            df = self.unified_df.copy()
            if target_sparte:
                df = df[df["Sparte"] == target_sparte]

            target_risk = None
            if any(x in ql for x in ["hoch", "high", "kritisch", "critical", "gefahr", "danger"]):
                target_risk = "Hoch"
            elif any(x in ql for x in ["mittel", "medium"]):
                target_risk = "Mittel"
            elif any(x in ql for x in ["niedrig", "low", "gut", "good"]):
                target_risk = "Niedrig"

            if target_risk:
                filtered = df[df["Risiko"] == target_risk]
            else:
                filtered = df

            total = len(df)
            lines = [
                f"📊 **Risiko-Übersicht{' — ' + target_sparte if target_sparte else ''}**",
                f"Gesamt: {total} Anschlüsse analysiert.",
                "",
            ]

            # Risk count summary
            risk_counts = df["Risiko"].value_counts()
            for level, emoji in [("Hoch", "🔴"), ("Mittel", "🟡"), ("Niedrig", "🟢"), ("Unbekannt", "⚪")]:
                cnt = risk_counts.get(level, 0)
                pct = round(100 * cnt / max(total, 1), 1)
                lines.append(f"{emoji} **{level}**: {cnt} ({pct}%)")

            lines.append("")
            lines.append("**Wie wird das Risiko berechnet?**")
            lines.append(
                "Das Risiko basiert auf dem Alter der Leitung im Vergleich zu den "
                "technischen Nutzungsdauern je Material (DVGW-Richtlinien):"
            )

            # Show thresholds per sparte
            sparten_to_show = [target_sparte] if target_sparte else list(RISK_LADDER.keys())
            for sp in sparten_to_show:
                if sp not in RISK_LADDER: continue
                lines.append(f"\n*{sp}:*")
                seen = set()
                for p in RISK_LADDER[sp]:
                    if p["material"] in seen: continue
                    seen.add(p["material"])
                    lines.append(
                        f"  • {p['material']}: Niedrig ≤{p['gut']}J | "
                        f"Mittel {p['gut']+1}–{p['mittel']}J | Hoch >{p['mittel']}J"
                    )

            # If a specific risk level was requested, show top examples
            if target_risk and not filtered.empty:
                lines.append(f"\n**Beispiele mit Risiko '{target_risk}'** (Top 5):")
                for _, r in filtered.sort_values("Alter", ascending=False).head(5).iterrows():
                    lines.append(
                        f"  → {r.get('Sparte','')} | {r.get('Straße','?')} {r.get('Hausnummer','')} | "
                        f"Material: {r.get('Werkstoff','?')} | Alter: {int(r.get('Alter',0))}J"
                    )

            return {
                "answer": "\n".join(lines),
                "hits": [], "model_used": "Risk-Overview-Engine-v2", "switched": True
            }

        # ── 4. Count Queries ──────────────────────────────────────────────────
        if any(x in ql for x in ["wie viele", "how many", "how much", "anzahl", "count",
                                   "summe", "total", "wieviel", "kritisch", "critical"]):
            target_sparte = None
            if "gas" in ql:    target_sparte = "Gas"
            elif "wasser" in ql: target_sparte = "Wasser"
            elif "strom" in ql:  target_sparte = "Strom"

            target_risk = None
            if any(x in ql for x in ["risk", "risiko", "kritisch", "critical", "gefahr", "danger"]):
                if any(x in ql for x in ["hoch", "high", "kritisch", "critical"]):
                    target_risk = "Hoch"
                elif any(x in ql for x in ["mittel", "medium"]):
                    target_risk = "Mittel"
                elif any(x in ql for x in ["niedrig", "low"]):
                    target_risk = "Niedrig"
                else:
                    target_risk = "Hoch"  # Default for unspecified "at risk"

            df = self.unified_df.copy()
            if target_sparte:
                df = df[df["Sparte"] == target_sparte]
            if target_risk:
                df = df[df["Risiko"] == target_risk]

            total = len(df)
            if total >= 0:
                risk_str = f" mit Risiko-Status **{target_risk}**" if target_risk else ""
                sparte_str = target_sparte or "Netz"
                msg = f"Wir haben insgesamt **{total} {sparte_str}-Anschlüsse**{risk_str} im System erfasst."
                if not target_sparte and not target_risk:
                    counts = self.unified_df["Sparte"].value_counts().to_dict()
                    msg += " Aufgeschlüsselt nach Sparten: " + ", ".join([f"{k}: {v}" for k, v in counts.items()])
                if target_risk:
                    msg += f"\n\n💡 *Hinweis: '{target_risk}' bedeutet, die Leitung hat die technische Nutzungsdauer für ihr Material überschritten oder befindet sich im kritischen Bereich. Siehe Risiko-Tabelle für Details.*"
                return {"answer": msg, "hits": [], "model_used": "Count-Engine-v3", "switched": True}

        # ── 5. Greetings ──────────────────────────────────────────────────────
        if ql.strip() in ["hi", "hallo", "hello", "guten tag", "hey"]:
            return {
                "answer": (
                    "Guten Tag! Ich bin das ESC Infrastructure Intelligence System. "
                    "Wie kann ich Ihnen heute bei der Analyse Ihrer Gas- oder Wasser-Anschlussdaten helfen?"
                ),
                "hits": [], "model_used": "System", "switched": False
            }

        # ── 6. Analytical Trends ──────────────────────────────────────────────
        if any(x in ql for x in ["material", "werkstoff", "anschlussart"]) and \
           any(x in ql for x in ["wann", "verbaut", "historisch", "zeitraum"]):
            try:
                df = self.unified_df.copy()
                df["Dekade"] = (df["Einbaujahr"] // 10) * 10
                summary = df.groupby(["Dekade", "Werkstoff"]).size().unstack(fill_value=0)
                lines = ["📜 **Historische Analyse: Materialverwendung**", ""]
                for decade in sorted(summary.index.dropna()):
                    if decade < 1920: continue
                    mats = summary.loc[decade]
                    top = mats[mats > 0].sort_values(ascending=False)
                    mat_str = ", ".join([f"{m} ({q})" for m, q in top.items()])
                    lines.append(f"- **{int(decade)}er Jahre**: {mat_str}")
                return {
                    "answer": "\n".join(lines), "hits": [],
                    "model_used": "History-Engine", "switched": True,
                    "download_data": summary.to_csv().encode('utf-8-sig')
                }
            except:
                pass

        # ── 7. Numeric Age Filters (with street / Gemeinde / Sparte sub-filters) ──
        age_match = re.search(r'(älter|older|>|über|over|more\s*than)\s*(?:als\s*|than\s*)?(\d+)', ql)
        if age_match:
            try:
                threshold = int(age_match.group(2))
                df_af = self.unified_df.copy()

                # Sparte filter
                af_sparte = None
                if "gas" in ql:       af_sparte = "Gas"
                elif "wasser" in ql:  af_sparte = "Wasser"
                if af_sparte:
                    df_af = df_af[df_af["Sparte"] == af_sparte]

                # Street filter — compare query against actual street names
                af_street = None
                unique_streets = [str(s) for s in df_af["Straße"].dropna().unique() if len(str(s)) > 3]
                matched_streets = [s for s in unique_streets if s.lower() in ql]
                if matched_streets:
                    af_street = max(matched_streets, key=len)
                    df_af = df_af[df_af["Straße"] == af_street]

                # Gemeinde / Ortsteil filter
                af_gemeinde = None
                if "Gemeinde" in df_af.columns:
                    unique_gem = [str(g) for g in df_af["Gemeinde"].dropna().unique() if len(str(g)) > 3]
                    matched_gem = [g for g in unique_gem if g.lower() in ql]
                    if matched_gem:
                        af_gemeinde = max(matched_gem, key=len)
                        df_af = df_af[df_af["Gemeinde"] == af_gemeinde]

                matches = df_af[df_af["Alter"] > threshold].sort_values("Alter", ascending=False)

                location_label = ""
                if af_street:    location_label += f" in **{af_street}**"
                if af_gemeinde:  location_label += f" ({af_gemeinde})"
                sparte_label    = f" — {af_sparte}" if af_sparte else ""

                if not matches.empty:
                    avg_age = matches["Alter"].mean()
                    lines = [
                        f"📊 **{len(matches)} Anschlüsse{sparte_label} älter als {threshold} Jahre{location_label}**",
                        f"Ø Alter der Treffer: **{avg_age:.0f} Jahre** | Max: **{int(matches['Alter'].max())} Jahre**",
                        "",
                    ]
                    # Risk breakdown
                    risk_cnt = matches["Risiko"].value_counts()
                    for lvl, em in [("Hoch","🔴"),("Mittel","🟡"),("Niedrig","🟢")]:
                        if risk_cnt.get(lvl, 0):
                            lines.append(f"{em} {lvl}: {risk_cnt[lvl]}")
                    lines.append("\n**Top Treffer:**")
                    for _, r in matches.head(10).iterrows():
                        lines.append(
                            f"  → {r.get('Sparte','')} | **{r.get('Straße','?')} {r.get('Hausnummer','')}** | "
                            f"Alter: **{int(r['Alter'])}J** | Material: {r.get('Werkstoff','?')} | "
                            f"Risiko: {r.get('Risiko','?')}"
                        )
                    dl_requested = any(x in ql for x in ["excel","csv","tabelle","table","format","tabular","alle","all"])
                    resp = {"answer": "\n".join(lines), "hits": [], "model_used": "Age-Filter-Engine-v2", "switched": True}
                    if dl_requested:
                        resp["download_data"] = matches.to_csv(index=False).encode('utf-8-sig')
                    return resp
                else:
                    return {
                        "answer": f"📊 Keine Anschlüsse{sparte_label} älter als {threshold} Jahre{location_label} gefunden.",
                        "hits": [], "model_used": "Age-Filter-Engine-v2", "switched": True,
                    }
            except:
                pass

        # ── 8. Full Table / Listing ───────────────────────────────────────────
        table_keywords = ["liste", "tabelle", "alle", "list", "table", "all",
                          "übersicht", "total", "excel", "csv", "daten", "format"]
        if any(x in ql for x in table_keywords):
            target_sparte = None
            if "gas" in ql:    target_sparte = "Gas"
            elif "wasser" in ql: target_sparte = "Wasser"

            df = self.unified_df.copy()
            if target_sparte:
                df = df[df["Sparte"].astype(str).str.contains(target_sparte, case=False, na=False)]

            # Smart NLP Filters
            if any(x in ql for x in ["high risk", "hohes risiko", "risiko hoch", "hoch"]):
                df = df[df["Risiko"] == "Hoch"]
            elif any(x in ql for x in ["medium", "mittel", "mittleres risiko"]):
                df = df[df["Risiko"] == "Mittel"]
            elif any(x in ql for x in ["low", "niedrig", "gering"]):
                df = df[df["Risiko"] == "Niedrig"]

            unique_streets = [str(s) for s in df["Straße"].dropna().unique() if len(str(s)) > 3]
            matched_streets = [s for s in unique_streets if s.lower() in ql]
            if matched_streets:
                best_street = max(matched_streets, key=lambda x: len(x))
                df = df[df["Straße"] == best_street]
                hn_match = re.search(r'\b\d{1,4}[a-zA-Z]?\b', ql.replace(best_street.lower(), ''))
                if hn_match:
                    num = hn_match.group()
                    df = df[df["Hausnummer"].astype(str).str.lower() == num.lower()]

            if target_sparte:
                answer = f"✅ **{target_sparte}-Übersicht**: {len(df)} Anschlüsse gefunden."
            else:
                answer = f"✅ **Gesamtübersicht**: {len(df)} Anschlüsse gefunden."

            if not df.empty:
                cols = ["Kundenname", "Kundennummer", "Sparte", "Straße", "Hausnummer", "Risiko", "Alter"]
                available_cols = [c for c in cols if c in df.columns]
                preview_df = df[available_cols].head(15)
                table_md = preview_df.to_markdown(index=False)
                dl_requested = any(x in ql for x in ["excel", "csv", "tabelle", "table", "format", "tabular"])
                resp = {
                    "answer": (
                        f"{answer}\n\n{table_md}\n\n"
                        "*(Oben sehen Sie die ersten 15 Zeilen. "
                        "Nutzen Sie den Button unten für den vollständigen Export als CSV)*"
                    ),
                    "hits": [], "model_used": "Full-Table-Engine", "switched": True
                }
                if dl_requested:
                    resp["download_data"] = df.to_csv(index=False).encode('utf-8-sig')
                return resp
            else:
                return {
                    "answer": (
                        f"📊 **Tabelle**: Keine Daten für "
                        f"'{target_sparte or 'Gesamt'}' gefunden, die Ihrer Anfrage entsprechen."
                    ),
                    "hits": [], "model_used": "Full-Table-Engine", "switched": True
                }

        # ── 9. General Map Navigation ─────────────────────────────────────────
        map_keywords = ["karte", "map", "landkarte", "netz-karte", "zeige", "view", "show", "öffne", "open"]
        target_sparte = None
        if "gas" in ql:    target_sparte = "Gas"
        elif "wasser" in ql: target_sparte = "Wasser"

        if any(x in ql for x in map_keywords) and not id_match and not any(x in ql for x in table_keywords):
            return {
                "answer": "🗺️ **Ich erstelle die Netz-Karte direkt hier im Chat...**",
                "hits": [], "model_used": "Navigation-Engine-v1", "switched": True,
                "pending_action": {
                    "type": "navigate_map_general",
                    "args": {"filter": target_sparte or "All"}
                }
            }

        # ── 10. Missing / Incomplete Documents ────────────────────────────────
        doc_kw = ["fehl", "dokument", "plan", "protokoll", "akte", "abnahme",
                  "vollständig", "lückenhaft", "missing", "unvollständig",
                  "materialangab", "unterlage", "inspektion"]
        if any(x in ql for x in doc_kw):
            df_doc = self.unified_df.copy()
            doc_sparte = None
            if "gas" in ql:    doc_sparte = "Gas"
            elif "wasser" in ql: doc_sparte = "Wasser"
            if doc_sparte:
                df_doc = df_doc[df_doc["Sparte"] == doc_sparte]

            total_doc = len(df_doc)
            missing_doc    = df_doc[df_doc["Dokumente"] == "Lückenhaft"] if "Dokumente" in df_doc.columns else pd.DataFrame()
            no_inspection  = df_doc[df_doc["missing_inspection"] == True] if "missing_inspection" in df_doc.columns else pd.DataFrame()

            lines = [
                f"📋 **Dokumenten-Analyse{' — ' + doc_sparte if doc_sparte else ' (Alle Sparten)'}**",
                f"Anschlüsse gesamt: **{total_doc}**",
                "",
                f"🔴 Fehlende Vertragsunterlagen (Gestattung/Auftrag): **{len(missing_doc)}**"
                f" ({round(100*len(missing_doc)/max(total_doc,1),1)}%)",
                f"⚠️  Kein Inspektionsdatum hinterlegt: **{len(no_inspection)}**"
                f" ({round(100*len(no_inspection)/max(total_doc,1),1)}%)",
                "",
                "**Geprüfte Dokumentenfelder:**",
                "  • Gestattungsvertrag (Grundstückseigentümer-Zustimmung)",
                "  • Auftragsunterlagen (Installateur)",
                "  • Anfragedokumentation",
                "  • (Letztes) Inspektionsdatum",
                "  *(Lageplan, Abnahmeprotokoll, Materialangabe: aus Primärsystem ergänzen)*",
            ]
            if not missing_doc.empty:
                lines.append("\n**Anschlüsse mit Dokumentenlücken (Top 5):**")
                for _, r in missing_doc.sort_values("Alter", ascending=False).head(5).iterrows():
                    lines.append(
                        f"  → {r.get('Sparte','')} | {r.get('Straße','?')} {r.get('Hausnummer','')} | "
                        f"Alter: {int(r.get('Alter',0))}J | Risiko: {r.get('Risiko','?')}"
                    )
            resp = {"answer": "\n".join(lines), "hits": [], "model_used": "Docs-Engine-v1", "switched": True}
            if any(x in ql for x in ["csv","excel","tabelle","format","tabular","export"]):
                combined_doc = pd.concat([missing_doc, no_inspection]).drop_duplicates()
                if not combined_doc.empty:
                    resp["download_data"] = combined_doc.to_csv(index=False).encode('utf-8-sig')
            return resp

        # ── 11. Renewal Due Analysis ──────────────────────────────────────────
        renewal_kw = ["erneuert werden", "erneuert", "erneuerung", "fällig",
                      "renewal", "ablauf", "nächsten jahren", "überfällig", "erneuern",
                      "sollten erneuert", "due for renewal"]
        horizon_match = re.search(r'(?:nächsten?\s+)?(\d+)\s*(?:jahre|years|j\.?)', ql)
        year_horizon  = int(horizon_match.group(1)) if horizon_match else None

        if any(x in ql for x in renewal_kw) and not any(x in ql for x in ["warum","why","grund","erklär"]):
            df_ren = self.unified_df.copy()
            ren_sparte = None
            if "gas" in ql:    ren_sparte = "Gas"
            elif "wasser" in ql: ren_sparte = "Wasser"
            if ren_sparte:
                df_ren = df_ren[df_ren["Sparte"] == ren_sparte]

            if "Erneuerung_empfohlen_bis" not in df_ren.columns:
                return {"answer": "Keine Erneuerungsdaten verfügbar.", "hits": [], "model_used": "Renewal-Engine-v1", "switched": True}

            df_ren = df_ren.copy()
            df_ren["Erneuerung_empfohlen_bis"] = pd.to_numeric(df_ren["Erneuerung_empfohlen_bis"], errors="coerce")
            df_ren = df_ren.dropna(subset=["Erneuerung_empfohlen_bis"])

            horizon = year_horizon if year_horizon else 10
            now     = CURRENT_YEAR
            overdue = df_ren[df_ren["Erneuerung_empfohlen_bis"] < now]
            due_now = df_ren[
                (df_ren["Erneuerung_empfohlen_bis"] >= now) &
                (df_ren["Erneuerung_empfohlen_bis"] <= now + horizon)
            ]

            lines = [
                f"🔄 **Erneuerungs-Analyse{' — ' + ren_sparte if ren_sparte else ''}**",
                "",
                f"🔴 **Bereits überfällig** (Erneuerungsjahr < {now}): **{len(overdue)} Anschlüsse**",
                f"⚠️  **Fällig in den nächsten {horizon} Jahren** (bis {now+horizon}): **{len(due_now)} Anschlüsse**",
                "",
            ]

            # Material breakdown for overdue
            if not overdue.empty:
                mat_cnt = overdue["Werkstoff"].value_counts()
                lines.append("**Überfällig — Aufschlüsselung nach Material:**")
                for mat, cnt in mat_cnt.head(5).items():
                    lines.append(f"  • {mat}: {cnt} Anschlüsse")
                lines.append("")

            # Dringendste Beispiele
            combined_due = pd.concat([overdue, due_now]).sort_values("Erneuerung_empfohlen_bis")
            if not combined_due.empty:
                lines.append("**Dringendste Erneuerungen (Top 5):**")
                for _, r in combined_due.head(5).iterrows():
                    yr = int(r["Erneuerung_empfohlen_bis"])
                    lines.append(
                        f"  → {r.get('Sparte','')} | {r.get('Straße','?')} {r.get('Hausnummer','')} | "
                        f"Material: {r.get('Werkstoff','?')} | Alter: {int(r.get('Alter',0))}J | "
                        f"Erneuerung bis: **{yr}**"
                    )

            resp = {"answer": "\n".join(lines), "hits": [], "model_used": "Renewal-Engine-v1", "switched": True}
            if any(x in ql for x in ["csv","excel","tabelle","format","tabular","export","alle","all"]):
                resp["download_data"] = combined_due.to_csv(index=False).encode('utf-8-sig')
            return resp

        # ── 12. Bundling / Koordinierungsanalyse ─────────────────────────────
        bundle_kw = ["bündeln", "bundle", "zusammenfassen", "straßenzug",
                     "kombinier", "bündelung", "batch", "gemeinsam erneuern",
                     "koordinier", "gemeinsame erneuerung"]
        if any(x in ql for x in bundle_kw):
            df_bun = self.unified_df.copy()
            bun_sparte = None
            if "gas" in ql:    bun_sparte = "Gas"
            elif "wasser" in ql: bun_sparte = "Wasser"
            if bun_sparte:
                df_bun = df_bun[df_bun["Sparte"] == bun_sparte]

            # Restrict to connections due within 15 years for relevance
            if "Erneuerung_empfohlen_bis" in df_bun.columns:
                df_bun["Erneuerung_empfohlen_bis"] = pd.to_numeric(df_bun["Erneuerung_empfohlen_bis"], errors="coerce")
                df_bun = df_bun[df_bun["Erneuerung_empfohlen_bis"] <= CURRENT_YEAR + 15]

            lines = [
                f"🔗 **Bündelungs-Analyse{' — ' + bun_sparte if bun_sparte else ''}**",
                "*Identifiziert Anschlüsse, die sich wirtschaftlich gemeinsam erneuern lassen.*",
                "",
            ]

            # Group by Straße
            if "Straße" in df_bun.columns and not df_bun.empty:
                sg = df_bun.groupby("Straße").agg(
                    Anzahl=("Kundennummer", "count"),
                    Frueheste_Erneuerung=("Erneuerung_empfohlen_bis", "min"),
                ).sort_values("Anzahl", ascending=False)
                top_streets = sg[sg["Anzahl"] >= 2].head(8)
                if not top_streets.empty:
                    lines.append("**Straßenzüge mit mehreren fälligen Anschlüssen:**")
                    for street, row in top_streets.iterrows():
                        yr_str = str(int(row["Frueheste_Erneuerung"])) if pd.notna(row["Frueheste_Erneuerung"]) else "?"
                        lines.append(
                            f"  🏘️ **{street}**: {int(row['Anzahl'])} Anschlüsse | früheste Erneuerung: {yr_str}"
                        )

            # Group by Dekade + Material
            if "Einbaujahr" in df_bun.columns and not df_bun.empty:
                df_bun["Dekade"] = (df_bun["Einbaujahr"] // 10) * 10
                dg = (
                    df_bun.groupby(["Dekade","Werkstoff"])
                    .agg(Anzahl=("Kundennummer","count"))
                    .reset_index()
                    .sort_values("Anzahl", ascending=False)
                )
                lines.append("\n**Bündelung nach Baujahr-Dekade & Material:**")
                for _, row in dg[dg["Anzahl"] >= 2].head(6).iterrows():
                    if pd.notna(row["Dekade"]):
                        lines.append(f"  📅 {int(row['Dekade'])}er — {row['Werkstoff']}: {int(row['Anzahl'])} Anschlüsse")

            lines.append("\n💡 *Gleichzeitige Baustellen auf demselben Straßenzug senken Tiefbau- und Koordinierungskosten um 20–35%.*")
            resp = {"answer": "\n".join(lines), "hits": [], "model_used": "Bundle-Engine-v1", "switched": True}
            if any(x in ql for x in ["csv","excel","tabelle","format","tabular","export"]):
                resp["download_data"] = df_bun.to_csv(index=False).encode('utf-8-sig')
            return resp

        # ── 13. Heat Pump / Charging Infrastructure Suitability ────────────────
        infra_kw = ["wärmepumpe", "heat pump", "ladeinfrastruktur", "ladesäule",
                    "e-auto", "elektromobil", "eignung", "ungeeignet für",
                    "geeignet für", "modernisierung infrastruktur"]
        if any(x in ql for x in infra_kw):
            df_inf = self.unified_df.copy()
            inf_sparte = None
            if "gas" in ql:    inf_sparte = "Gas"
            elif "wasser" in ql: inf_sparte = "Wasser"
            if inf_sparte:
                df_inf = df_inf[df_inf["Sparte"] == inf_sparte]

            is_hp      = any(x in ql for x in ["wärmepumpe","heat pump","heizung"])
            is_charge  = any(x in ql for x in ["lade","e-auto","elektro","charging"])

            lines = []
            if is_hp or not is_charge:
                gas_df = df_inf[df_inf["Sparte"] == "Gas"] if "Sparte" in df_inf.columns else df_inf
                unsuitable = gas_df[gas_df["Risiko"].isin(["Hoch","Mittel"])] if not gas_df.empty else pd.DataFrame()
                over_life  = gas_df[gas_df.get("Infrastruktur_ungeeignet", pd.Series(False, index=gas_df.index)) == True] if not gas_df.empty else pd.DataFrame()

                lines += [
                    "🌡️ **Wärmepumpen-Eignung — Gasanschlüsse**",
                    "",
                    "*(Annahme: Bei Umstieg auf Wärmepumpe entfällt Gasbedarf.*",
                    " *Gasanschlüsse mit Risiko Hoch/Mittel oder abgelaufener Lebensdauer sollten*",
                    " *im Zuge des Umstiegs stillgelegt oder ersetzt werden.)*",
                    "",
                    f"⚠️  Gasanschlüsse mit Risiko Hoch/Mittel: **{len(unsuitable)}**",
                    f"🔴 Lebensdauer bereits überschritten: **{len(over_life)}**",
                ]
                if not unsuitable.empty:
                    mat_cnt = unsuitable["Werkstoff"].value_counts()
                    lines.append("  Materialien: " + ", ".join([f"{m} ({c})" for m, c in mat_cnt.head(4).items()]))

            if is_charge or not is_hp:
                lines += [
                    "",
                    "⚡ **Ladeinfrastruktur (E-Mobilität)**",
                    "*(Hinweis: Kapazitäts- und Absicherungsdaten der Stromanschlüsse sind*",
                    " *für diese Prüfung erforderlich. In der aktuellen Gas/Wasser-Datenbasis*",
                    " *nicht enthalten — bitte Stromdaten gesondert einpflegen.)*",
                ]

            return {"answer": "\n".join(lines), "hits": [], "model_used": "InfraSuitability-Engine-v1", "switched": True}

        # ── 14. Muster / Pattern Analysis ─────────────────────────────────────
        pattern_kw = ["muster", "pattern", "auffällig", "erkenntnisse", "was fällt",
                      "gemeinsamkeit", "was ist auffällig", "überblick akte",
                      "was zeigt", "daten auswerten"]
        if any(x in ql for x in pattern_kw):
            df_pat = self.unified_df.copy()
            pat_sparte = None
            if "gas" in ql:    pat_sparte = "Gas"
            elif "wasser" in ql: pat_sparte = "Wasser"
            if pat_sparte:
                df_pat = df_pat[df_pat["Sparte"] == pat_sparte]

            lines = [
                f"🔎 **Muster-Analyse der Hausanschlussakten{' — ' + pat_sparte if pat_sparte else ''}**",
                "",
            ]

            # Age structure
            if "Alter" in df_pat.columns and not df_pat["Alter"].isna().all():
                lines.append(
                    f"📅 **Altersstruktur**: Ø {df_pat['Alter'].mean():.0f} J | "
                    f"Median {df_pat['Alter'].median():.0f} J | Max {df_pat['Alter'].max():.0f} J"
                )

            # Dominant materials
            if "Werkstoff" in df_pat.columns:
                top_m = df_pat["Werkstoff"].value_counts().head(4)
                lines.append("🏗️ **Dominante Materialien**: " + ", ".join([f"{m} ({c})" for m, c in top_m.items()]))

            # Risk distribution
            if "Risiko" in df_pat.columns:
                rc = df_pat["Risiko"].value_counts()
                total_p = len(df_pat)
                hoch_pct  = round(100 * rc.get("Hoch",0)  / max(total_p,1), 1)
                mittel_pct= round(100 * rc.get("Mittel",0) / max(total_p,1), 1)
                lines.append(
                    f"⚠️  **Risiko**: 🔴 Hoch {hoch_pct}% ({rc.get('Hoch',0)}) | "
                    f"🟡 Mittel {mittel_pct}% ({rc.get('Mittel',0)}) | "
                    f"🟢 Niedrig {round(100*rc.get('Niedrig',0)/max(total_p,1),1)}% ({rc.get('Niedrig',0)})"
                )

            # Document gaps
            if "Dokumente" in df_pat.columns:
                miss_d = (df_pat["Dokumente"] == "Lückenhaft").sum()
                lines.append(f"📋 **Dokumentenlücken**: {miss_d} ({round(100*miss_d/max(len(df_pat),1),1)}%) unvollständig")

            # Missing inspections
            if "missing_inspection" in df_pat.columns:
                no_insp = df_pat["missing_inspection"].sum()
                lines.append(f"🔍 **Kein Inspektionsdatum**: {no_insp} ({round(100*no_insp/max(len(df_pat),1),1)}%)")

            # Material trend by decade
            if "Einbaujahr" in df_pat.columns and "Werkstoff" in df_pat.columns:
                df_pat["Dekade"] = (df_pat["Einbaujahr"] // 10) * 10
                dom = df_pat.groupby("Dekade")["Werkstoff"].agg(lambda x: x.mode()[0] if len(x) else "n/a")
                recent = dom[dom.index >= 1960]
                if not recent.empty:
                    lines.append("\n**Dominantes Material je Jahrzehnt:**")
                    for decade, mat in recent.items():
                        if pd.notna(decade):
                            lines.append(f"  • {int(decade)}er: {mat}")

            return {"answer": "\n".join(lines), "hits": [], "model_used": "Pattern-Engine-v1", "switched": True}

        # ── 15. Störungs-Korrelation / Failure Risk ────────────────────────────
        corr_kw = ["zusammenhang", "korrelation", "störung", "ausfall",
                   "störungswahrscheinlich", "schadenswahrscheinlich", "failure",
                   "correlation", "ausfallrisiko", "störungsrate"]
        if any(x in ql for x in corr_kw):
            df_cor = self.unified_df.copy()
            cor_sparte = None
            if "gas" in ql:    cor_sparte = "Gas"
            elif "wasser" in ql: cor_sparte = "Wasser"
            if cor_sparte:
                df_cor = df_cor[df_cor["Sparte"] == cor_sparte]

            lines = [
                f"📊 **Korrelations-Analyse: Baujahr × Material × Störungsrisiko**",
                "*(Annahme Testbetrieb: Störungsraten nach DVGW-Richtwerten.*",
                " *Echte Störungshistorie aus IMS/GIS ergänzt Präzision deutlich.)*",
                "",
                "**Angenommene Jahres-Störungsrate nach Risikostufe (DVGW):**",
                "  🔴 Hoch  → 8–15% pro Jahr",
                "  🟡 Mittel → 2–5%  pro Jahr",
                "  🟢 Niedrig → < 1% pro Jahr",
                "",
            ]

            # Material × Risk correlation
            if "Werkstoff" in df_cor.columns and "Risiko" in df_cor.columns:
                mr = df_cor.groupby("Werkstoff")["Risiko"].value_counts().unstack(fill_value=0)
                if "Hoch" in mr.columns:
                    mr["Hoch_Pct"] = (mr["Hoch"] / mr.sum(axis=1) * 100).round(1)
                    top5 = mr.sort_values("Hoch_Pct", ascending=False).head(5)
                    lines.append("**Materialien nach Hoch-Risiko-Anteil:**")
                    for mat, row in top5.iterrows():
                        total_m = int(row.get("Hoch",0) + row.get("Mittel",0) + row.get("Niedrig",0))
                        lines.append(
                            f"  • {mat}: **{row['Hoch_Pct']}%** Hoch-Risiko "
                            f"({int(row.get('Hoch',0))}/{total_m})"
                        )

            # Age brackets correlation
            if "Alter" in df_cor.columns:
                bins = [0, 20, 40, 60, 80, 200]
                labels = ["0–20J","21–40J","41–60J","61–80J",">80J"]
                df_cor["Altersgruppe"] = pd.cut(df_cor["Alter"], bins=bins, labels=labels, right=True)
                age_risk = df_cor.groupby("Altersgruppe")["Risiko"].value_counts().unstack(fill_value=0)
                if "Hoch" in age_risk.columns:
                    lines.append("\n**Hoch-Risiko-Anteil nach Altersgruppe:**")
                    for ag, row in age_risk.iterrows():
                        total_ag = int(row.sum())
                        pct = round(100*row.get("Hoch",0)/max(total_ag,1),1)
                        lines.append(f"  • {ag}: {pct}% Hoch ({int(row.get('Hoch',0))}/{total_ag})")

            lines.append("\n💡 *Für genaue Prognose: Störungsdaten aus Betriebsdatenbank importieren.*")
            return {"answer": "\n".join(lines), "hits": [], "model_used": "Correlation-Engine-v1", "switched": True}

        return None

    # ─────────────────────────────────────────────────────────────────────────
    # MAIN ANSWER ENTRYPOINT
    # FIX: LLM system prompt now injects FULL engineering standards (no 1500-char
    #      truncation). Risk questions that fall through to the LLM will still
    #      get accurate threshold context.
    # ─────────────────────────────────────────────────────────────────────────
    def answer_question(    
        self,
        question: str,
        utility: Optional[str] = None,
        history: Optional[List[Dict]] = None
    ) -> Dict[str, Any]:
        ql = (question or "").lower()

        update_keywords = ["update", "ändern", "setze", "change", "aktualisier",
                           "korrigier", "fix", "put", "schreib"]
        is_update = any(x in ql for x in update_keywords)
        ql = (question or "").lower()

# Stronger update detection
        update_keywords = [
    "update", "ändern", "setze", "change", "fix", "modify",
    "aktualisier", "korrigier", "put", "schreib", "replace"
        ]

# detect presence of ID (very important)
        has_id = bool(re.search(r'\b\d+\b', ql))

# detect assignment patterns
        assignment_patterns = [
            r"\bset\b.*\bto\b",        # set X to Y
            r"\bchange\b.*\bto\b",     # change X to Y
            r"\bupdate\b",             # update
            r"\bmodify\b",             # modify
        ]

        has_assignment = any(re.search(p, ql) for p in assignment_patterns)

        is_update = (
            any(x in ql for x in update_keywords)
            and has_id
            and has_assignment
)



        route = self._route_query(question)  

        # Build context-enriched query for follow-up questions
        search_query = question
        if history:
            history_to_use = (
                history[:-1]
                if history and history[-1].get("content", "") == question
                else history
            )
            last_user_msg = next(
                (h["content"] for h in reversed(history_to_use) if h["role"] == "user"), ""
            )
            if last_user_msg and len(question.split()) < 10 and not re.search(r'\d+', question):
                search_query = f"{last_user_msg}. {question}"

        route = self._route_query(question)

        # ALWAYS try dataframe first UNLESS it's an update
        if not is_update:
            df_res = self._try_dataframe_answer(search_query)
            if df_res:
                return df_res
            
        if OFFLINE:
            return {
                "answer": (
                    "KI-Assistent ist im Offline-Modus. "
                    "Bitte stellen Sie eine Internetverbindung her, um AI-Antworten zu erhalten."
                ),
                "hits": [], "model_used": "Offline-Mode", "switched": False
            }

        status = self.check_llm_status()
        if not status["ok"]:
            return {
                "answer": f"⚠️ **Service-Status**: {status['msg']}",
                "hits": [], "model_used": "Status-Check", "switched": False
            }

        hits = []
        try:
            results = self.vs.query(
                query_embeddings=self.embedder.embed([search_query]), top_k=4
            )
            hits = [
                {"meta": m, "doc": d, "score": 1 - dist}
                for m, d, dist in zip(
                    results["metadatas"][0],
                    results["documents"][0],
                    results["distances"][0]
                )
            ]
            ctx = "\n---\n".join([h["doc"] for h in hits])

            if self.azure_openai_endpoint:
                headers = {
                    "api-key": self.llm_api_key,
                    "Content-Type": "application/json"
                }
                url = self.llm_base_url
            else:
                headers = {
                    "Authorization": f"Bearer {self.llm_api_key}",
                    "Content-Type": "application/json"
                }
                url = f"{self.llm_base_url}/chat/completions"

            # ── AGENTIC MODE (write operations) ──────────────────────────────
            if is_update:
                tools = [
                    {
                        "type": "function",
                        "function": {
                            "name": "update_asset",
                            "description": (
                                "Updates any field of a utility asset in the database. "
                                "Call this when the user asks to change, set, or update any value."
                            ),
                            "parameters": {
                                "type": "object",
                                "properties": {
                                    "customer_id": {
                                        "type": "string",
                                        "description": "Customer ID, e.g. '3' or 'Kunde 3'."
                                    },
                                    "field_name": {
                                        "type": "string",
                                        "description": (
                                            "Column to update, e.g. Hausnummer, Schutzrohr, "
                                            "Werkstoff, Installateur Name, Gemeinde, "
                                            "Gestattungsvertrag, etc."
                                        )
                                    },
                                    "new_value": {
                                        "type": "string",
                                        "description": "New value to write."
                                    },
                                    "utility": {
                                        "type": "string",
                                        "enum": ["Gas", "Wasser", "Gemeinsam"],
                                        "description": "Utility sector. Use Gemeinsam for shared fields."
                                    }
                                },
                                "required": ["customer_id", "field_name", "new_value", "utility"]
                            }
                        }
                    },
                    {
                        "type": "function",
                        "function": {
                            "name": "navigate_to_map",
                            "description": (
                                "Navigates to the map view for a specific customer or the general map. "
                                "Call this when the user asks to see a customer, an address, or the map."
                            ),
                            "parameters": {
                                "type": "object",
                                "properties": {
                                    "customer_id": {
                                        "type": "string",
                                        "description": "Optional Customer ID. If omitted, shows general map."
                                    }
                                }
                            }
                        }
                    }
                ]
                # FIX: Full engineering standards injected into agentic mode too
                system_content = (
                    """
You are the ESC Infrastructure Intelligence Assistant.

STRICT RULES:

- Only use provided data
- Never guess
- If missing → say "Keine Daten vorhanden"
- Keep answers short and numeric

RISK RULES:
Always include:
- Material
- Age
- Threshold

Format:
Risiko: <LEVEL> — Material: <X>, Alter: <Y>J  
Schwelle: ≤A (Niedrig), B–C (Mittel), >C (Hoch)

If high → mention exceeded years

COUNT RULES:
- Always exact numbers
- No approximation

STYLE:
- Short
- Clean
- No explanation unless needed
"""  + "\n\n" + self.reference_manual
                )
                payload = {
                    "model": self.llm_model,
                    "messages": [{"role": "system", "content": system_content}],
                    "tools": tools,
                    "tool_choice": "auto",
                    "temperature": 0.0,
                    "max_tokens": 400
                }
                if history:
                    h_slice = history[:-1] if history[-1].get("content", "") == question else history
                    for h in h_slice[-6:]:
                        role = "assistant" if h["role"] == "bot" else h["role"]
                        payload["messages"].append({"role": role, "content": h["content"]})
                payload["messages"].append({"role": "user", "content": f"Question: {question}"})

            # ── FAST READ MODE ────────────────────────────────────────────────
            else:
                # FIX: Inject FULL engineering standards + full stats summary.
                # Old code did reference_manual[:1500] which truncated everything.
                system_content = (
                    "You are an infrastructure data expert for energy networks (Gas, Wasser, Strom).\n\n"
                    "Answer questions accurately based on the data and Engineering Standards below.\n"
                    "When explaining Risiko (Risk), you MUST reference the exact thresholds from the "
                    "Technical Life Tables — never guess or generalize.\n\n"
                    f"{self.reference_manual}\n"
                    "\nIMPORTANT:\n"
                    "- NEVER say you lack data for counts; use the summary above.\n"
                    "- For risk questions always cite material + age + threshold.\n"
                    "- Respond in the same language as the question."
                )
                payload = {
                    "model": self.llm_model,
                    "messages": [{"role": "system", "content": system_content}],
                    "temperature": 0.1,
                    "max_tokens": 800   # FIX: was 600, increased for detailed risk explanations
                }
                if history:
                    h_slice = history[:-1] if history[-1].get("content", "") == question else history
                    for h in h_slice[-6:]:
                        role = "assistant" if h["role"] == "bot" else h["role"]
                        payload["messages"].append({"role": role, "content": h["content"]})
                payload["messages"].append({
                    "role": "user",
                    "content": f"Relevant data from database:\n{ctx}\n\nQuestion: {question}"
                })

            resp = requests.post(
                url,
                headers=headers, json=payload, timeout=30
            )
            if resp.status_code == 200:
                choice = resp.json()["choices"][0]
                msg = choice["message"]

                # Tool calls (agentic mode)
                if "tool_calls" in msg and msg["tool_calls"]:
                    tc   = msg["tool_calls"][0]["function"]
                    args = json.loads(tc["arguments"])

                    if tc["name"] == "update_asset":
                        result = self.apply_update(
                            customer_id=args.get("customer_id"),
                            field_name=args.get("field_name"),
                            new_value=args.get("new_value"),
                            utility=args.get("utility"),
                        )

                        return {
                            "answer": (
                                f"🛠️ **Update ausgeführt**\n\n{result.get('msg')}"
                                if result["ok"]
                                else f"❌ Fehler beim Update:\n{result.get('msg')}"
                            ),
                            "hits": hits,
                            "model_used": self.llm_model,
                            "pending_action": None
                        }




                    if tc["name"] == "navigate_to_map":
                        sid = args.get("customer_id")
                        if not sid:
                            return {
                                "answer": "🗺️ **Ich öffne die Netz-Karte für Sie...**",
                                "hits": hits, "model_used": self.llm_model,
                                "pending_action": {"type": "navigate_map_general", "args": {}}
                            }
                        if self.unified_df.empty:
                            self.unified_df = get_unified_df()
                        matches = pd.DataFrame()
                        for col in ["Kundennummer", "Objekt_ID", "Objekt-ID", "Objekt-ID_Global"]:
                            if col in self.unified_df.columns:
                                mask = self.unified_df[col].astype(str).str.contains(
                                    rf'\b{sid}\b', regex=True, na=False
                                )
                                if mask.any():
                                    matches = self.unified_df[mask]
                                    break
                        if not matches.empty:
                            res = matches.iloc[0]
                            if pd.notna(res.get("lat")) and pd.notna(res.get("lon")):
                                return {
                                    "answer": (
                                        f"📍 **Navigation zur Karte...**\n"
                                        f"Ich zeige Ihnen den Anschluss `{res.get('Kundennummer')}` auf der Karte."
                                    ),
                                    "hits": hits, "model_used": self.llm_model,
                                    "pending_action": {
                                        "type": "navigate_map",
                                        "args": {
                                            "customer_id": str(res.get("Kundennummer")),
                                            "lat": float(res["lat"]),
                                            "lon": float(res["lon"])
                                        }
                                    }
                                }
                        return {
                            "answer": "🗺️ **Ich öffne die Netz-Karte für Sie...**",
                            "hits": hits, "model_used": self.llm_model,
                            "pending_action": {"type": "navigate_map_general", "args": {}}
                        }

                answer = msg.get("content", "")
                return {"answer": answer, "hits": hits, "model_used": self.llm_model, "switched": False}

        except Exception:
            pass

        if hits:
            fallback = (
                "⏱️ **Zeitüberschreitung**: Relevantes aus der Datenbank:\n\n"
                + "\n".join([f"- {h['doc']}" for h in hits[:2]])
            )
            return {"answer": fallback, "hits": hits, "model_used": "Timeout-Fallback", "switched": False}

        return {"answer": "Keine Antwort gefunden.", "hits": [], "model_used": "Error", "switched": False}

    def stream_answer(
        self,
        question: str,
        utility: Optional[str] = None,
        history: Optional[List[Dict]] = None
    ):
        """
        Generator that yields SSE-formatted strings.
        Fast (dataframe/update) answers are emitted as a single 'done' event.
        LLM read-mode answers stream one token at a time.
        """
        ql = (question or "").lower()
        update_keywords = ["update", "ändern", "setze", "change", "aktualisier",
                           "korrigier", "fix", "put", "schreib"]
        is_update = any(x in ql for x in update_keywords)

        # Build enriched search query from history context (same as answer_question)
        search_query = question
        if history:
            history_to_use = (
                history[:-1] if history and history[-1].get("content", "") == question
                else history
            )
            last_user_msg = next(
                (h["content"] for h in reversed(history_to_use) if h["role"] == "user"), ""
            )
            if last_user_msg and len(question.split()) < 10 and not re.search(r'\d+', question):
                search_query = f"{last_user_msg}. {question}"

        # 1. Fast dataframe engine — no LLM, yield instantly as a single done event
        if not is_update:
            df_res = self._try_dataframe_answer(search_query)
            if df_res:
                yield (
                    f"data: {json.dumps({'type': 'done', 'answer': df_res.get('answer', ''), 'pending_action': df_res.get('pending_action')})}\n\n"
                )
                return

        # 2. Update commands (tool calls) and offline mode — non-streaming, single event
        if is_update or OFFLINE:
            result = self.answer_question(question, utility=utility, history=history)
            yield (
                f"data: {json.dumps({'type': 'done', 'answer': result.get('answer', ''), 'pending_action': result.get('pending_action')})}\n\n"
            )
            return

        # 3. LLM read-mode — stream tokens
        status = self.check_llm_status()
        if not status["ok"]:
            _msg = status["msg"]
            yield f"data: {json.dumps({'type': 'done', 'answer': f'⚠️ Service: {_msg}'})}\n\n"
            return

        try:
            results = self.vs.query(
                query_embeddings=self.embedder.embed([search_query]), top_k=4
            )
            ctx = "\n---\n".join(results["documents"][0])

            if self.azure_openai_endpoint:
                headers = {
                    "api-key": self.llm_api_key,
                    "Content-Type": "application/json",
                }
                url = self.llm_base_url
            else:
                headers = {
                    "Authorization": f"Bearer {self.llm_api_key}",
                    "Content-Type": "application/json",
                }
                url = f"{self.llm_base_url}/chat/completions"
            system_content = (
                "You are the ESC Infrastructure Intelligence System for gas and water utility networks.\n\n"

                "═══ STRICT RULES ═══\n"
                "1. ONLY use the provided data and Engineering Standards below\n"
                "2. NEVER invent counts, years, materials, or risk levels\n"
                "3. If data is missing → say exactly: 'Keine Daten vorhanden'\n"
                "4. Use ONLY the exact numbers from the CURRENT INFRASTRUCTURE SUMMARY for counts\n"
                "5. Respond in the same language as the user (German if asked in German)\n\n"

                "═══ RISK CALCULATION (DVGW) ═══\n"
                "Risk = f(Material, Age). Always cite:\n"
                "  • Material name\n"
                "  • Current age in years\n"
                "  • Threshold: Niedrig ≤ A J | Mittel A+1–B J | Hoch > B J\n\n"

                "═══ TRAINED QUESTION CATEGORIES ═══\n"
                "Answer these question types precisely using the infrastructure summary:\n\n"

                "1. ALTER-FRAGEN ('Welche Anschlüsse sind älter als X Jahre?')\n"
                "   → Filter by Alter > X. Group by Straße or Gemeinde if mentioned.\n"
                "   → Always state: total count, avg age, risk breakdown.\n\n"

                "2. DOKUMENTEN-FRAGEN ('Fehlen Dokumente / Pläne / Abnahmeprotokolle?')\n"
                "   → Use 'MISSING DOCUMENTS' from summary. Report: Lückenhaft count + inspection gaps.\n"
                "   → Note: Lageplan, Abnahmeprotokoll, Materialangabe → from Primärsystem.\n\n"

                "3. MATERIAL-ZEITRAUM ('Welche Materialien wurden wann verbaut?')\n"
                "   → Group by Einbaujahr-Dekade × Werkstoff. Show dominant material per decade.\n\n"

                "4. RISIKO-BEWERTUNG ('Welche Anschlüsse haben erhöhtes Risiko?')\n"
                "   → Use 'RISK BY SPARTE' from summary. Show material + age + threshold.\n\n"

                "5. BAUJAHR-MATERIAL-KORRELATION ('Zusammenhang Baujahr, Material, Störung?')\n"
                "   → DVGW-Annahme: Hoch=8–15%/J, Mittel=2–5%/J, Niedrig=<1%/J Störungsrate.\n"
                "   → Show material Hoch-Risiko percentages. Note: real data from IMS/GIS improves precision.\n\n"

                "6. ERNEUERUNGS-EMPFEHLUNG ('Welche Anschlüsse sollten erneuert werden?')\n"
                "   → Use 'RENEWAL OVERDUE' and 'RENEWAL DUE WITHIN 10 YEARS' from summary.\n"
                "   → Break down by material. List most urgent cases.\n\n"

                "7. BÜNDELUNG ('Welche Erneuerungen lassen sich bündeln?')\n"
                "   → Group by Straße (same street = shared trench). Group by Einbaujahr-Dekade (same material era).\n"
                "   → Tip: Koordinierte Baustellen senken Kosten 20–35%.\n\n"

                "8. INFRASTRUKTUR-EIGNUNG ('Geeignet für Wärmepumpen / Ladeinfrastruktur?')\n"
                "   → Wärmepumpen: Gas connections with Risiko Hoch/Mittel → decommissioning candidates.\n"
                "   → Ladeinfrastruktur: Requires Strom capacity data — not in current Gas/Wasser dataset.\n\n"

                "9. KOMBINIERTE FRAGEN → Answer all relevant sub-categories in sequence.\n\n"

                "10. MUSTER ('Welche Muster fallen in den Akten auf?')\n"
                "    → Report: avg age, dominant materials, risk%, document gaps%, inspection gaps%, decade trends.\n\n"

                f"{self.reference_manual}\n\n"

                "IMPORTANT: For all count-based answers use ONLY the numbers in the summary above.\n"
                "Never approximate. Never guess. If uncertain → say 'Bitte konkreten Anschluss oder Sparte angeben.'\n"
            )
            payload = {
                "model": self.llm_model,
                "messages": [{"role": "system", "content": system_content}],
                "temperature": 0.1,
                "max_tokens": 800,
                "stream": True,
            }
            if history:
                h_slice = history[:-1] if history[-1].get("content", "") == question else history
                for h in h_slice[-6:]:
                    role = "assistant" if h["role"] == "bot" else h["role"]
                    payload["messages"].append({"role": role, "content": h["content"]})
            payload["messages"].append({
                "role": "user",
                "content": f"Relevant data from database:\n{ctx}\n\nQuestion: {question}",
            })

            resp = requests.post(
                url,
                headers=headers, json=payload, timeout=60, stream=True
            )
            if resp.status_code != 200:
                yield f"data: {json.dumps({'type': 'done', 'answer': f'⚠️ LLM Error {resp.status_code}'})}\n\n"
                return

            for raw_line in resp.iter_lines():
                if not raw_line:
                    continue
                line = raw_line.decode("utf-8") if isinstance(raw_line, bytes) else raw_line
                if not line.startswith("data: "):
                    continue
                data_str = line[6:].strip()
                if data_str == "[DONE]":
                    yield f"data: {json.dumps({'type': 'done'})}\n\n"
                    return
                try:
                    chunk = json.loads(data_str)
                    token = chunk["choices"][0]["delta"].get("content", "")
                    if token:
                        yield f"data: {json.dumps({'type': 'token', 'token': token})}\n\n"
                except Exception:
                    pass

            yield f"data: {json.dumps({'type': 'done'})}\n\n"

        except Exception as e:
            _err = str(e)[:120]
            yield f"data: {json.dumps({'type': 'done', 'answer': f'⚠️ Fehler: {_err}'})}\n\n"

    def chat_general(self, user_message: str, history: List[Dict]) -> Dict[str, Any]:
        return self.answer_question(user_message, history=history)